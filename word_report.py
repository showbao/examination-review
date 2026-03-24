import re
from io import BytesIO
from collections import Counter
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

from utils import (
    clean_markdown_symbol,
    is_markdown_table_line,
    is_markdown_separator_line,
    parse_markdown_table_rows,
    normalize_analysis_tables,
    _normalize_table_header,
)

# ==========================================
# 字型設定
# ==========================================

def set_font_style(run, size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

# ==========================================
# 標題識別與正規化
# ==========================================

def is_main_section_header(line):
    clean = clean_markdown_symbol(line.strip())
    main_headers = [
        "總結與建議", "題幹與邏輯品質", "素養導向深度審查",
        "公平性與敏感度審查", "雙向細目表核算",
        "難易度與負擔分析", "評量診斷與補救教學"
    ]
    return clean in main_headers

SUB_HEADER_ALIASES = {
    "最優先修正": "最優先修正",
    "最優先修正 (Critical)": "最優先修正",
    "難度與鑑別度點評": "難度與鑑別度點評",
    "值得讚許之處": "值得讚許之處",
    "後續優化建議": "後續優化建議",
    "優良試題": "優良試題",
    "待確認試題及修改建議": "待確認試題及修改建議",
    "真素養": "真素養",
    "假素養/待確認": "假素養/待確認",
    "通過 (無敏感議題)": "通過 (無敏感議題)",
    "潛在爭議 (具體指出問題)": "潛在爭議 (具體指出問題)",
    "整體作答負擔觀察": "整體作答負擔觀察",
    "閱讀負擔": "閱讀負擔",
    "圖表判讀負擔": "圖表判讀負擔",
    "運算負擔": "運算負擔"
}

TEXT_MODE_SUBHEADERS = {"閱讀負擔", "圖表判讀負擔", "運算負擔"}

def normalize_sub_header_text(text):
    text = clean_markdown_symbol(text.strip())
    text = re.sub(r'^[✏️📖📊🧮🔴⚖️👍💡🟢🟡⚠️📌]+\s*', '', text)
    text = re.sub(r'\s*[:：]\s*$', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def canonical_sub_header(text):
    normalized = normalize_sub_header_text(text)
    if normalized.startswith("最優先修正"):
        return "最優先修正"
    if normalized.startswith("難度與鑑別度點評"):
        return "難度與鑑別度點評"
    if normalized.startswith("值得讚許之處"):
        return "值得讚許之處"
    if normalized.startswith("後續優化建議"):
        return "後續優化建議"
    if normalized.startswith("待確認試題及修改建議"):
        return "待確認試題及修改建議"
    for alias, canonical in SUB_HEADER_ALIASES.items():
        if normalized.startswith(alias):
            return canonical
    return normalized

def is_sub_section_header(line):
    normalized = normalize_sub_header_text(line)
    allowed_headers = set(SUB_HEADER_ALIASES.values())
    for header in allowed_headers:
        if re.fullmatch(rf'{re.escape(header)}\s*[：:]?\s*', normalized):
            return True
    return False

def normalize_compare_text(text):
    text = clean_markdown_symbol(text.strip())
    text = re.sub(r'\s+', '', text)
    text = re.sub(r'[：:，,。．、；;（）()]', '', text)
    return text.strip()

# ==========================================
# Word 段落輸出函式
# ==========================================

def add_main_section_title(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font_style(run, size=16, bold=True, color=RGBColor(91, 124, 153))
    run.font.underline = True

def add_sub_section_title(container, text):
    spacer = container.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(canonical_sub_header(text))
    set_font_style(run, size=13, bold=True, color=RGBColor(111, 133, 158))
    return p

def add_bullet_to_cell(container, text):
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(0)
    run_bullet = p.add_run("• ")
    set_font_style(run_bullet, size=12)
    run_text = p.add_run(text)
    set_font_style(run_text, size=12)
    return p

def add_indented_sub_bullet_to_cell(container, text):
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(2.0)
    p.paragraph_format.first_line_indent = Cm(0)
    run_bullet = p.add_run("－ ")
    set_font_style(run_bullet, size=12, color=RGBColor(111, 133, 158))
    label_match = re.match(r'^(【[^】]+】：?)(.*)', text, re.DOTALL)
    if label_match:
        run_label = p.add_run(label_match.group(1))
        set_font_style(run_label, size=12, bold=True)
        run_content = p.add_run(label_match.group(2))
        set_font_style(run_content, size=12)
    else:
        run_text = p.add_run(text)
        set_font_style(run_text, size=12)
    return p

def add_plain_text_to_cell(container, text):
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_font_style(run, size=12)
    return p

def add_note_text(container, text):
    """在表格下方渲染註解文字（灰色小字）"""
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    set_font_style(run, size=10, color=RGBColor(120, 120, 120))

# ==========================================
# 表格輔助
# ==========================================

def set_cell_shading(cell, fill="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def prevent_row_break(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)

# 判斷一行是否為表格註解行的 pattern
_NOTE_PATTERN = re.compile(
    r'^\s*[\(（]?\s*(?:註|Note|備註|說明|因|＊|※)',
    re.IGNORECASE,
)


def _is_table_note_row(cells):
    """判斷某一列的內容是否為註解（如「（註：因版面缺失...）」）"""
    # 如果這一列只有第一個 cell 有文字，且符合註解 pattern
    combined = " ".join(c.strip() for c in cells).strip()
    if not combined:
        return False
    # 全部文字合併檢查
    return bool(_NOTE_PATTERN.match(combined))


def _render_word_table(container, data):
    """渲染 Markdown 表格為 Word 表格，若最後幾列為註解則抽出到表格外"""
    if not data:
        return
    rows_data = parse_markdown_table_rows(data) if isinstance(data[0], str) else data
    if not rows_data:
        return

    # 分離註解列：從尾端開始檢查
    note_rows = []
    while len(rows_data) > 1 and _is_table_note_row(rows_data[-1]):
        note_rows.insert(0, rows_data.pop())

    if not rows_data:
        return

    rows = len(rows_data)
    cols = max(len(row) for row in rows_data)

    table = container.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_texts = [str(x).strip() for x in rows_data[0]]
    if header_texts == ["認知向度", "對應題號", "比重"]:
        width_map = [Cm(4.0), Cm(10.0), Cm(4.0)]
    elif header_texts == ["難易度", "對應題號", "比重"]:
        width_map = [Cm(4.0), Cm(10.0), Cm(4.0)]
    else:
        width_map = [Cm(6.0)] * cols

    for r in range(rows):
        row = table.rows[r]
        prevent_row_break(row)
        for c in range(cols):
            cell = table.cell(r, c)
            text = rows_data[r][c] if c < len(rows_data[r]) else ""
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c < len(width_map):
                cell.width = width_map[c]
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = True if r < rows - 1 else False
                for run in p.runs:
                    set_font_style(run, size=12)
            if r == 0:
                set_cell_shading(cell, "D9D9D9")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True

    # 渲染抽出的註解列到表格下方
    for note_row in note_rows:
        note_text = " ".join(c.strip() for c in note_row if c.strip())
        if note_text:
            add_note_text(container, note_text)

# ==========================================
# Word 表格去重
# ==========================================

def _deduplicate_word_table_lines(lines):
    table_blocks = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if is_markdown_table_line(stripped):
            start = i
            first_line = stripped
            while i < len(lines) and (is_markdown_table_line(lines[i].strip()) or is_markdown_separator_line(lines[i].strip())):
                i += 1
            table_blocks.append({
                "start": start, "end": i,
                "header_norm": _normalize_table_header(first_line),
            })
        else:
            i += 1
    if not table_blocks:
        return lines
    header_counts = Counter(tb["header_norm"] for tb in table_blocks)
    duplicate_headers = {h for h, c in header_counts.items() if c > 1}
    if not duplicate_headers:
        return lines
    last_occurrence = {}
    for tb in table_blocks:
        if tb["header_norm"] in duplicate_headers:
            last_occurrence[tb["header_norm"]] = tb["start"]
    remove_lines = set()
    for tb in table_blocks:
        if tb["header_norm"] in duplicate_headers and tb["start"] != last_occurrence[tb["header_norm"]]:
            for j in range(tb["start"], tb["end"]):
                remove_lines.add(j)
    return [lines[i] for i in range(len(lines)) if i not in remove_lines]


# ==========================================
# Word 報告主函式
# ==========================================

def create_word_report(analysis_text, metadata):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    style.font.size = Pt(12)

    # 標題
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("臺中市北屯區建功國小評量AI審題報告")
    set_font_style(run, size=18, bold=True)
    doc.add_paragraph()

    # 卷頭資訊表格
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("學年度：", metadata.get("year", "   "), "學期：", metadata.get("semester", "   ")),
        ("年級：", metadata.get("grade", "   "), "科目：", metadata.get("subject", "   ")),
        ("評量類別：", metadata.get("exam_type", "   "), "審查日期：", time.strftime("%Y/%m/%d")),
        ("命題者簽名：", "          ", "審題者簽名：", "          ")
    ]
    for row_idx, row_data in enumerate(info_data):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_font_style(run, size=12, bold=(col_idx % 2 == 0))
    doc.add_paragraph()

    # 整合通知區塊
    notice_lines = []
    curriculum_display = metadata.get("curriculum_display", "")
    if curriculum_display:
        notice_lines.append(curriculum_display)
    notice_lines.append(
        "⚠️ 系統限制：本系統僅針對試題內容進行深度分析，"
        "未檢核試卷形式（如題號連貫性、配分加總正確性），請老師務必自行審閱。"
    )
    notice_table = doc.add_table(rows=1, cols=1)
    notice_table.style = 'Table Grid'
    notice_cell = notice_table.cell(0, 0)
    for existing_p in notice_cell.paragraphs:
        existing_p.text = ""
    for idx, line in enumerate(notice_lines):
        if idx == 0:
            p_notice = notice_cell.paragraphs[0]
        else:
            p_notice = notice_cell.add_paragraph()
        run = p_notice.add_run(line)
        if line.startswith("⚠️"):
            set_font_style(run, size=11, bold=True, color=RGBColor(200, 80, 80))
        else:
            set_font_style(run, size=11, bold=False)
    doc.add_paragraph()

    # 命題教師修改及說明
    p = doc.add_paragraph()
    run = p.add_run("命題教師修改及說明")
    set_font_style(run, size=14, bold=True, color=RGBColor(91, 124, 153))

    feedback_table = doc.add_table(rows=1, cols=1)
    feedback_table.style = 'Table Grid'
    cell = feedback_table.cell(0, 0)
    checkbox_items = [
        "無需修改試題。",
        "經命題老師確認，以下試題為AI幻覺，已判斷無需修正。\n   試題：_______________________________________________________",
        "經命題老師確認，以下試題已修正。\n   試題：_______________________________________________________"
    ]
    for item_text in checkbox_items:
        p_check = cell.add_paragraph()
        p_check.paragraph_format.line_spacing = 2.0
        run_box = p_check.add_run("□ ")
        set_font_style(run_box, size=18)
        run_text = p_check.add_run(item_text)
        set_font_style(run_text, size=12)

    p_other = cell.add_paragraph("其他說明：")
    set_font_style(p_other.runs[0] if p_other.runs else p_other.add_run("其他說明："), size=12)
    p_other.paragraph_format.line_spacing = 2.0
    for _ in range(5):
        p_empty = cell.add_paragraph()
        p_empty.paragraph_format.line_spacing = 2.0
    doc.add_paragraph()

    doc.add_page_break()

    # --- 內容解析主迴圈 ---
    analysis_text = normalize_analysis_tables(analysis_text)
    lines = analysis_text.split('\n')
    lines = _deduplicate_word_table_lines(lines)

    table_mode = False
    table_data = []
    current_sub_header = None

    for raw_line in lines:
        stripped_line = raw_line.strip()

        if not stripped_line:
            if table_mode and table_data:
                _render_word_table(doc, table_data)
                table_mode = False
                table_data = []
            continue

        clean_text = clean_markdown_symbol(stripped_line)
        clean_text = re.sub(r'\s*[:：]\s*$', '', clean_text)

        if is_main_section_header(stripped_line):
            if table_mode and table_data:
                _render_word_table(doc, table_data)
                table_mode = False
                table_data = []
            current_sub_header = None
            add_main_section_title(doc, clean_text)
            continue

        if is_sub_section_header(stripped_line):
            if table_mode and table_data:
                _render_word_table(doc, table_data)
                table_mode = False
                table_data = []
            matched_sub_header = canonical_sub_header(stripped_line)
            add_sub_section_title(doc, matched_sub_header)
            current_sub_header = matched_sub_header
            continue

        if is_markdown_table_line(stripped_line):
            table_mode = True
            table_data.append(stripped_line)
            continue

        if table_mode and table_data:
            _render_word_table(doc, table_data)
            table_mode = False
            table_data = []

        if not clean_text:
            continue

        clean_text = re.sub(r'^\d+\.\s*', '', clean_text)

        if current_sub_header in TEXT_MODE_SUBHEADERS:
            normalized_clean = normalize_compare_text(clean_text)
            normalized_header = normalize_compare_text(current_sub_header)
            if normalized_clean.startswith(normalized_header):
                clean_text = re.sub(
                    rf'^\s*{re.escape(current_sub_header)}\s*[：:，,。．、；;（）()]*\s*',
                    '', clean_text
                ).strip()
                if not clean_text:
                    continue
            clean_text = re.sub(r'^[•\*\-]\s*', '', clean_text).strip()
            if clean_text:
                add_plain_text_to_cell(doc, clean_text)
            continue

        if current_sub_header is not None:
            normalized_clean = normalize_compare_text(clean_text)
            normalized_header = normalize_compare_text(current_sub_header)
            if normalized_clean.startswith(normalized_header):
                clean_text = re.sub(
                    rf'^\s*{re.escape(current_sub_header)}\s*[：:，,。．、；;（）()]*\s*',
                    '', clean_text
                ).strip()
                if not clean_text:
                    continue

        if re.match(r'^[\*\-]\s+', stripped_line) or re.match(r'^\d+\.\s*', stripped_line):
            if (current_sub_header == "待確認試題及修改建議"
                    and re.match(r'^【[^】]+】', clean_text)):
                add_indented_sub_bullet_to_cell(doc, clean_text)
            else:
                add_bullet_to_cell(doc, clean_text)
        else:
            if current_sub_header is not None:
                add_bullet_to_cell(doc, clean_text)
            else:
                add_plain_text_to_cell(doc, clean_text)

    if table_mode and table_data:
        _render_word_table(doc, table_data)

    doc.add_page_break()

    # 評量診斷與補救教學
    add_main_section_title(doc, "評量診斷與補救教學")
    remedial_table = doc.add_table(rows=1, cols=1)
    remedial_table.style = 'Table Grid'
    remedial_cell = remedial_table.cell(0, 0)
    remedial_lines = [
        "僅針對錯誤率較高的題目（一至二題）進行試題分析／學習診斷。",
        "",
        "【題目】 第（   ）大題第（   ）題",
        "",
        "【學習表現／學習內容】",
        "",
        "",
        "【評量結果分析】（為何此題錯誤率高？可從試題內容、教學方法等層面分析）",
        "",
        "",
        "",
        "",
        "【可行補救教學策略】（請列舉具體可行之策略）",
        "",
        "",
        "",
        "",
    ]
    for line in remedial_lines:
        p_rem = remedial_cell.add_paragraph()
        run = p_rem.add_run(line)
        set_font_style(run, size=12)

    f = BytesIO()
    doc.save(f)
    return f.getvalue()
