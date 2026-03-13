import streamlit as st
import google.generativeai as genai
import os
import time
import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

# ==========================================
# 0. 視覺風格設定 (莫蘭迪色系 & CSS)
# ==========================================
st.set_page_config(page_title="北屯區建功國小AI審題系統", page_icon="📝", layout="wide")

morandi_css = """
<style>
    /* 全站字體放大 */
    html, body, [class*="css"] {
        font-size: 20px; 
    }
    
    .stApp { background-color: #F5F7F7; }
    
    /* 標題層級與大小調整 */
    h1 { color: #5B7C99 !important; font-family: 'Helvetica Neue', sans-serif; font-size: 2.5rem !important; }
    h2 { color: #5B7C99 !important; font-family: 'Helvetica Neue', sans-serif; font-size: 2.2rem !important; font-weight: bold !important; }
    h3 { color: #5B7C99 !important; font-family: 'Helvetica Neue', sans-serif; font-size: 1.3rem !important; font-weight: normal !important; }
    
    /* 針對總結區塊的特別樣式 */
    .summary-box {
        background-color: #E3F2FD;
        border-left: 5px solid #2196F3;
        padding: 20px;
        border-radius: 5px;
        margin-bottom: 20px;
    }

    /* 按鈕樣式：簡約細框 + 自適應大小 */
    div.stButton > button {
        background-color: #FFFFFF !important;      /* 內部留白 (無填滿) */
        color: #5B7C99 !important;                 /* 文字顏色 (莫蘭迪藍灰) */
        border: 1px solid #E0E0E0 !important;      /* 極細灰框線 */
        border-radius: 12px;                       /* 圓角 */
        
        /* 核心設計：灰白陰影 (創造浮起感) */
        box-shadow: 3px 3px 8px rgba(0, 0, 0, 0.05), -2px -2px 6px #FFFFFF !important;
        
        /* 改為自適應大小 */
        height: auto !important;
        width: auto !important;
        padding: 10px 25px !important;
        margin-top: 10px;
        
        font-weight: bold;
        font-size: 1.1rem;
        
        /* 動畫過渡 */
        transition: all 0.3s ease;
        display: flex;
        align-items: center;
        justify-content: center;
    }

    /* 滑鼠懸停效果 (Hover) */
    div.stButton > button:hover {
        background-color: #FDFDFD !important;
        color: #8DA399 !important;
        border: 1px solid #8DA399 !important;
        transform: translateY(-2px);
        box-shadow: 5px 5px 12px rgba(0, 0, 0, 0.1), -3px -3px 8px #FFFFFF !important;
    }
    
    /* 資訊看板樣式放大 */
    .dashboard-card {
        background-color: #E8ECEC; padding: 20px; border-radius: 10px; border-left: 6px solid #8DA399; 
        margin-bottom: 25px; color: #4A4A4A;
        font-size: 1.1rem; 
    }
    
    .footer {
        position: fixed; left: 0; bottom: 0; width: 100%; background-color: #F5F7F7; color: #888; text-align: center; padding: 15px; font-size: 14px; border-top: 1px solid #ddd; z-index: 999;
    }
    .footer-spacer { height: 60px; }
</style>
"""
st.markdown(morandi_css, unsafe_allow_html=True)

# ==========================================
# 1. 系統全域設定 (Global Config)
# ==========================================

ENABLE_ADVANCED_FEATURES = False 

# ==========================================
# 2. 輔助函式：模型管理與 Word 生成
# ==========================================

def get_best_flash_model(api_key):
    genai.configure(api_key=api_key)
    try:
        models = [m for m in genai.list_models() if 'generateContent' in m.supported_generation_methods and "flash" in m.name.lower() and "gemini" in m.name.lower()]
        models.sort(key=lambda x: x.name, reverse=True)
        return models[0].name if models else "models/gemini-1.5-flash"
    except: return "models/gemini-1.5-flash"

def get_best_pro_model(api_key):
    genai.configure(api_key=api_key)
    try:
        models = [m for m in genai.list_models() if 'generateContent' in m.supported_generation_methods and "pro" in m.name.lower() and "gemini" in m.name.lower()]
        models.sort(key=lambda x: x.name, reverse=True)
        return models[0].name if models else "models/gemini-1.5-pro"
    except: return "models/gemini-1.5-pro"

def upload_to_gemini(file_obj):
    import tempfile
    suffix = ".pdf" if file_obj.name.endswith(".pdf") else ".jpg"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_obj.getvalue())
        tmp_path = tmp.name
    file_ref = genai.upload_file(tmp_path, mime_type="application/pdf" if suffix == ".pdf" else "image/jpeg")
    while file_ref.state.name == "PROCESSING":
        time.sleep(1)
        file_ref = genai.get_file(file_ref.name)
    os.remove(tmp_path)
    return file_ref

# --- Word 生成核心邏輯 ---
def set_font_style(run, size=12, bold=False, color=None):
    """設定字體為標楷體 + Times New Roman"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def clean_markdown_symbol(text):
    text = text.replace("**", "").replace("__", "")
    text = text.replace("<br>", "").replace("<br/>", "").replace("<br />", "")
    
    # 移除標題符號 #
    text = re.sub(r'#+\s*', '', text) 
    
    # 移除開頭列表符號
    text = re.sub(r'^[\*\-]\s+', '', text)
   
    # 移除燈號
    for icon in ["🟢", "🔴", "🟡", "⚠️", "👍", "💡", "📊", "⚖️", "📖", "📊", "🧮"]:
        text = text.replace(icon, "")
    
    text = text.lstrip("*- ")
    return text.strip()

def is_markdown_table_line(line):
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 3

def is_markdown_separator_line(line):
    compact = line.strip().replace(" ", "")
    return bool(re.fullmatch(r'\|?[:\-\|]+\|?', compact)) and "-" in compact

def parse_markdown_table_rows(table_lines):
    rows = []
    max_cols = 0

    for raw in table_lines:
        if is_markdown_separator_line(raw):
            continue

        cells = [clean_markdown_symbol(c.strip()) for c in raw.strip().strip("|").split("|")]
        rows.append(cells)
        max_cols = max(max_cols, len(cells))

    for row in rows:
        if len(row) < max_cols:
            row.extend([""] * (max_cols - len(row)))

    return rows

def normalize_analysis_tables(text):
    lines = text.split("\n")
    normalized = []
    i = 0

    while i < len(lines):
        if is_markdown_table_line(lines[i]):
            block = []
            while i < len(lines) and is_markdown_table_line(lines[i]):
                block.append(lines[i])
                i += 1

            rows = parse_markdown_table_rows(block)
            if rows:
                header = rows[0]
                col_count = len(header)

                normalized.append("| " + " | ".join(header) + " |")
                normalized.append("| " + " | ".join(["---"] * col_count) + " |")

                for row in rows[1:]:
                    normalized.append("| " + " | ".join(row[:col_count]) + " |")
            continue

        normalized.append(lines[i])
        i += 1

    return "\n".join(normalized)

def is_main_section_header(line):
    clean = clean_markdown_symbol(line.strip())
    main_headers = [
        "總結與建議",
        "題幹與邏輯品質",
        "素養導向深度審查",
        "公平性與敏感度審查",
        "雙向細目表核算",
        "難易度與負擔分析",
        "評量診斷與補救教學"
    ]
    return clean in main_headers

def is_sub_section_header(line):
    normalized = normalize_sub_header_text(line)

    allowed_headers = set(SUB_HEADER_ALIASES.values())

    for header in allowed_headers:
        if re.fullmatch(rf'{re.escape(header)}\s*[：:]?\s*', normalized):
            return True

    return False

def add_main_section_title(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font_style(run, size=16, bold=True, color=RGBColor(91, 124, 153))
    run.font.underline = True

def add_sub_section_title(container, text):
    spacer = container.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)   # 子標內縮
    run = p.add_run(canonical_sub_header(text))  # 不再顯示圖示
    set_font_style(run, size=13, bold=True, color=RGBColor(111, 133, 158))
    return p

def add_bullet_to_cell(container, text):
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(0)

    run_bullet = p.add_run("• ")
    set_font_style(run_bullet, size=12)

    run_text = p.add_run(text)
    set_font_style(run_text, size=12)
    return p

def add_plain_text_to_cell(container, text):
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_font_style(run, size=12)
    return p

SUB_HEADER_ICON_MAP = {
    "最優先修正": "🔴",
    "最優先修正 (Critical)": "🔴",
    "難度與鑑別度點評": "⚖️",
    "值得讚許之處": "👍",
    "後續優化建議": "💡",
    "優良試題": "🟢",
    "待確認試題": "🟡",
    "真素養": "🟢",
    "假素養/待確認": "🟡",
    "通過 (無敏感議題)": "🟢",
    "潛在爭議 (具體指出問題)": "⚠️",
    "整體作答負擔觀察": "📌",
    "閱讀負擔": "📖",
    "圖表判讀負擔": "📊",
    "運算負擔": "🧮"
}

SUB_HEADER_ALIASES = {
    "最優先修正": "最優先修正",
    "最優先修正 (Critical)": "最優先修正",
    "難度與鑑別度點評": "難度與鑑別度點評",
    "值得讚許之處": "值得讚許之處",
    "後續優化建議": "後續優化建議",
    "優良試題": "優良試題",
    "待確認試題": "待確認試題",
    "真素養": "真素養",
    "假素養/待確認": "假素養/待確認",
    "通過 (無敏感議題)": "通過 (無敏感議題)",
    "潛在爭議 (具體指出問題)": "潛在爭議 (具體指出問題)",
    "整體作答負擔觀察": "整體作答負擔觀察",
    "閱讀負擔": "閱讀負擔",
    "圖表判讀負擔": "圖表判讀負擔",
    "運算負擔": "運算負擔"
}

def normalize_sub_header_text(text):
    text = clean_markdown_symbol(text.strip())
    text = re.sub(r'^[📖📊🧮🔴⚖️👍💡🟢🟡⚠️📌]+\s*', '', text)
    text = re.sub(r'\s*[:：]\s*$', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def canonical_sub_header(text):
    normalized = normalize_sub_header_text(text)

    # 優先處理容易變形的標題
    if normalized.startswith("最優先修正"):
        return "最優先修正"
    if normalized.startswith("難度與鑑別度點評"):
        return "難度與鑑別度點評"
    if normalized.startswith("值得讚許之處"):
        return "值得讚許之處"
    if normalized.startswith("後續優化建議"):
        return "後續優化建議"

    for alias, canonical in SUB_HEADER_ALIASES.items():
        if normalized.startswith(alias):
            return canonical
    return normalized
    
def normalize_compare_text(text):
    text = clean_markdown_symbol(text.strip())
    text = re.sub(r'\s+', '', text)
    text = re.sub(r'[：:，,。．、；;（）()]', '', text)
    return text.strip()

TEXT_MODE_SUBHEADERS = {
    "閱讀負擔",
    "圖表判讀負擔",
    "運算負擔"
}    
def decorate_sub_header(text):
    canonical = canonical_sub_header(text)
    icon = SUB_HEADER_ICON_MAP.get(canonical, "")
    return f"{icon} {canonical}".strip()

def set_cell_shading(cell, fill="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def prevent_row_break(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split) 

def create_word_report(analysis_text, metadata):
    doc = Document()
    
    # 設定邊界 (1.27 cm)
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # 預設樣式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    style.font.size = Pt(12)

    # 標題
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("臺中市北屯區建功國小評量AI審題報告")
    set_font_style(run, size=18, bold=True)
    doc.add_paragraph()

    # 卷頭資訊
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("學年度：", metadata.get("year", "   "), "學期：", metadata.get("semester", "   ")),
        ("年級：", metadata.get("grade", "   "), "科目：", metadata.get("subject", "   ")),
        ("評量類別：", metadata.get("exam_type", "   "), "審查日期：", time.strftime("%Y/%m/%d")),
        ("命題者簽名：", "          ", "審題者簽名：", "          ")
    ]
    for row_idx, row_data in enumerate(info_data):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_font_style(run, size=12, bold=(col_idx % 2 == 0))
            
    doc.add_paragraph()

    # 命題教師修改及說明
    p = doc.add_paragraph()
    run = p.add_run("命題教師修改及說明")
    set_font_style(run, size=14, bold=True, color=RGBColor(91, 124, 153))
    
    feedback_table = doc.add_table(rows=1, cols=1)
    feedback_table.style = 'Table Grid'
    cell = feedback_table.cell(0, 0)
    
    checkbox_items = [
        "無需修改試題。",
        "經命題老師確認，以下試題為AI幻覺，已判斷無需修正。\n   試題：_______________________________________________________",
        "經命題老師確認，以下試題已修正。\n   試題：_______________________________________________________"
    ]
    
    for item_text in checkbox_items:
        p_check = cell.add_paragraph()
        p_check.paragraph_format.line_spacing = 2.0 
        
        run_box = p_check.add_run("□ ")
        set_font_style(run_box, size=18) 
        
        run_text = p_check.add_run(item_text)
        set_font_style(run_text, size=12)

    p_other = cell.add_paragraph("其他說明：")
    set_font_style(p_other.runs[0] if p_other.runs else p_other.add_run("其他說明："), size=12)
    p_other.paragraph_format.line_spacing = 2.0
    
    for _ in range(5):
        p_empty = cell.add_paragraph()
        p_empty.paragraph_format.line_spacing = 2.0
    
    doc.add_paragraph()

    # 系統免責聲明
    warning_text = "⚠️ 系統限制與聲明：本系統僅針對試題內容進行深度分析，未檢核「命題範圍」與「試卷形式」（如題號連貫性、配分加總正確性），請老師務必自行審閱。"
    p_warn = doc.add_paragraph()
    run_warn = p_warn.add_run(warning_text)
    set_font_style(run_warn, size=11, bold=True, color=RGBColor(255, 0, 0))

    doc.add_page_break()

    # --- 內容解析（主標題 / 子標題 / 外框區塊版）---
    analysis_text = normalize_analysis_tables(analysis_text)
    lines = analysis_text.split('\n')
    table_mode = False
    table_data = []

    # --- 內容解析（取消分析章節外框版）---
    analysis_text = normalize_analysis_tables(analysis_text)
    lines = analysis_text.split('\n')
    table_mode = False
    table_data = []

    current_sub_header = None
    current_main_section = None

    for raw_line in lines:
        stripped_line = raw_line.strip()

        if not stripped_line:
            if table_mode and table_data:
                _render_word_table(doc, table_data)
                table_mode = False
                table_data = []
            continue

        clean_text = clean_markdown_symbol(stripped_line)
        clean_text = re.sub(r'\s*[:：]\s*$', '', clean_text)

        if is_main_section_header(stripped_line):
            if table_mode and table_data:
                _render_word_table(doc, table_data)
                table_mode = False
                table_data = []

            current_main_section = clean_text
            current_sub_header = None
            add_main_section_title(doc, clean_text)
            continue

        if is_sub_section_header(stripped_line):
            if table_mode and table_data:
                _render_word_table(doc, table_data)
                table_mode = False
                table_data = []

            matched_sub_header = canonical_sub_header(stripped_line)

            add_sub_section_title(doc, matched_sub_header)
            current_sub_header = matched_sub_header
            continue

        if is_markdown_table_line(stripped_line):
            table_mode = True
            table_data.append(stripped_line)
            continue

        if table_mode and table_data:
            _render_word_table(doc, table_data)
            table_mode = False
            table_data = []

        if not clean_text:
            continue

        # 將數字列點統一轉成小黑點內容
        clean_text = re.sub(r'^\d+\.\s*', '', clean_text)

        # --- 副標文字模式：閱讀負擔 / 圖表判讀負擔 / 運算負擔 ---
        if current_sub_header in TEXT_MODE_SUBHEADERS:
            # 若內文又以目前副標開頭，去除重複標題字樣
            normalized_clean = normalize_compare_text(clean_text)
            normalized_header = normalize_compare_text(current_sub_header)

            if normalized_clean.startswith(normalized_header):
                clean_text = re.sub(
                    rf'^\s*{re.escape(current_sub_header)}\s*[：:，,。．、；;（）()]*\s*',
                    '',
                    clean_text
                ).strip()

                if not clean_text:
                    continue

            # 若 AI 仍以列點開頭，移除列點符號後當一般文字輸出
            clean_text = re.sub(r'^[•\*\-]\s*', '', clean_text).strip()

            if clean_text:
                add_plain_text_to_cell(doc, clean_text)
            continue

        # --- 一般副標 / 一般內容 ---
        if current_sub_header is not None:
            normalized_clean = normalize_compare_text(clean_text)
            normalized_header = normalize_compare_text(current_sub_header)

            if normalized_clean.startswith(normalized_header):
                clean_text = re.sub(
                    rf'^\s*{re.escape(current_sub_header)}\s*[：:，,。．、；;（）()]*\s*',
                    '',
                    clean_text
                ).strip()

                if not clean_text:
                    continue

        if re.match(r'^[\*\-]\s+', stripped_line) or re.match(r'^\d+\.\s*', stripped_line):
            add_bullet_to_cell(doc, clean_text)
        else:
            if current_sub_header is not None:
                add_bullet_to_cell(doc, clean_text)
            else:
                add_plain_text_to_cell(doc, clean_text)

    if table_mode and table_data:
        _render_word_table(doc, table_data)
        
    # 評量診斷與補救教學
    add_main_section_title(doc, "評量診斷與補救教學")

    remedial_table = doc.add_table(rows=1, cols=1)
    remedial_table.style = 'Table Grid'
    remedial_cell = remedial_table.cell(0, 0)

    remedial_lines = [
        "僅針對錯誤率較高的題目（一至二題）進行試題分析／學習診斷。",
        "",
        "【題目】 第（   ）大題第（   ）題",
        "",
        "【學習表現／學習內容】",
        "",
        "",
        "【評量結果分析】（為何此題錯誤率高？可從試題內容、教學方法等層面分析）",
        "",
        "",
        "",
        "",
        "【可行補救教學策略】（請列舉具體可行之策略）",
        "",
        "",
        "",
        "",
    ]

    for line in remedial_lines:
        p_rem = remedial_cell.add_paragraph()
        run = p_rem.add_run(line)
        set_font_style(run, size=12)

    from io import BytesIO
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def _render_word_table(container, data):
    if not data:
        return

    rows_data = parse_markdown_table_rows(data) if isinstance(data[0], str) else data
    if not rows_data:
        return

    rows = len(rows_data)
    cols = max(len(row) for row in rows_data)

    table = container.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_row = rows_data[0] if rows_data else []
    header_texts = [str(x).strip() for x in header_row]

    if header_texts == ["認知向度", "對應題號", "比重"]:
        width_map = [Cm(4.0), Cm(10.0), Cm(4.0)]
    elif header_texts == ["難易度", "對應題號", "比重"]:
        width_map = [Cm(4.0), Cm(10.0), Cm(4.0)]
    else:
        width_map = [Cm(6.0)] * cols

    for r in range(rows):
        row = table.rows[r]
        prevent_row_break(row)

        for c in range(cols):
            cell = table.cell(r, c)
            text = rows_data[r][c] if c < len(rows_data[r]) else ""
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if c < len(width_map):
                cell.width = width_map[c]

            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = True if r < rows - 1 else False

                for run in p.runs:
                    set_font_style(run, size=12)

            if r == 0:
                set_cell_shading(cell, "D9D9D9")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        
def clean_ai_hallucinations(text):
    """
    針對 AI 輸出進行清洗，但保護 Markdown 表格區塊
    """
    table_blocks = []

    def protect_table(match):
        table_blocks.append(match.group(0))
        return f"@@TABLE_BLOCK_{len(table_blocks)-1}@@"

    text = re.sub(r'((?:^\|.*\|\s*$\n?){2,})', protect_table, text, flags=re.MULTILINE)

    text = re.sub(r'^\s*[一二三四五六七八九十0-9]+[、\.][\u4e00-\u9fa5]+\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'([。；：])\s*([一二三四五六七八九十\d]+-[\d]+)', r'\1\n* \2', text)

    def truncate_list(match):
        header = match.group(1)
        content = match.group(2)
        lines = [line for line in content.split('\n') if line.strip()]

        if len(lines) > 5:
            new_content = "\n".join(lines[:5]) + "\n* (其餘優良試題略)..."
            return f"{header}\n{new_content}\n"
        return match.group(0)

    text = re.sub(r'(###\s*🟢\s*優良試題.*?)((\n\s*[\*\-].*?)+)(?=\n###)', truncate_list, text, flags=re.DOTALL)
    text = re.sub(r'(###\s*🟢\s*真素養.*?)((\n\s*[\*\-].*?)+)(?=\n###)', truncate_list, text, flags=re.DOTALL)
    text = re.sub(r'(###\s*🟢\s*通過.*?)((\n\s*[\*\-].*?)+)(?=\n###)', truncate_list, text, flags=re.DOTALL)

    text = re.sub(r'\n{3,}', '\n\n', text).strip()

    for idx, block in enumerate(table_blocks):
        text = text.replace(f"@@TABLE_BLOCK_{idx}@@", normalize_analysis_tables(block))

    return normalize_analysis_tables(text)

# ==========================================
# 3. 登入與介面
# ==========================================
def render_footer():
    st.markdown('<div class="footer-spacer"></div>', unsafe_allow_html=True)
    st.markdown('<div class="footer">Designed for 臺中市北屯區建功國小 | Powered by Gemini 3.0</div>', unsafe_allow_html=True)

def check_login():
    # 已登入時：直接放行，並把 email 存進 session_state 方便後續使用
    if st.user.is_logged_in:
        user_email = st.user.get("email", "")
        st.session_state["user_email"] = user_email
        return True

    # 未登入時：顯示聲明與 Google 登入按鈕
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.markdown("## 北屯區建功國小 AI 審題系統")
        st.markdown("### ⚠️ 使用前請務必詳閱免責聲明")
        st.markdown("""
        **使用前請詳閱以下說明：**
        1. **本系統運用 AI 技術輔助教師審閱試題，分析結果僅供教學參考。**
        2. **人工查核機制**：AI 生成內容可能存在誤差，最終試卷定稿請務必回歸教師專業判斷。**
        3. **資料隱私安全**：嚴禁上傳包含學生個資、隱私或機密敏感內容之文件。**
        4. **授權使用範圍**：本系統無償提供予臺中市北屯區建功國小教師使用。**
        """)

    btn_col1, btn_col2, btn_col3 = st.columns([1, 2, 1])
    with btn_col2:
        if st.button("我同意聲明並使用 Google 登入", use_container_width=True):
            st.login()
            st.stop()

    render_footer()
    return False

if not check_login():
    st.stop()

# ==========================================
# 4. 主流程與 Prompt
# ==========================================

if "analysis_result" not in st.session_state: st.session_state.analysis_result = None
if "used_model_name" not in st.session_state: st.session_state.used_model_name = ""
if "metadata" not in st.session_state: st.session_state.metadata = {}

st.title("北屯區建功國小AI審題系統")

user_email = st.session_state.get("user_email", "")
top_col1, top_col2 = st.columns([5, 1])
with top_col1:
    if user_email:
        st.caption(f"目前登入者：{user_email}")
with top_col2:
    if st.button("登出"):
        st.logout()

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    st.error("請設定 Secrets: GEMINI_API_KEY")
    st.stop()

# --- 介面佈局：上傳區 (單欄流式排版) ---
st.subheader(" 上傳試卷 ")
exam_file = st.file_uploader("📤 上傳試卷", type=["pdf", "jpg", "png"], key="exam_uploader", label_visibility="collapsed")

# --- 按鈕區 (位於上傳區正下方) ---
st.markdown('<div style="height: 15px;"></div>', unsafe_allow_html=True) 
start_btn = st.button(" 開始\n AI 審查", type="primary", use_container_width=True)

# --- 進階功能區 (預設隱藏) ---
context_files = None
unit_list = []
manual_model = "Gemini 3.0 Flash"
model_mode = "智慧分流"

if ENABLE_ADVANCED_FEATURES:
    st.markdown("---")
    st.markdown("#### ⚙️ 進階設定 (系統狀態、教材、單元、模型)")
    
    st.markdown("""
    <div class="dashboard-card">
        <b>⚪ 系統狀態：</b>待命中... 請上傳試卷<br>
        <small>啟用檔名路由：檔名含「數/理/化/生」➔ Pro | 其餘 ➔ Flash</small>
    </div>
    """, unsafe_allow_html=True)
    
    with st.expander("📂 參考教材與單元設定", expanded=False):
        col_adv_1, col_adv_2 = st.columns(2)
        with col_adv_1:
            st.markdown("**上傳課本、習作 (可選)**")
            context_files = st.file_uploader("拖曳參考教材", type=["pdf"], accept_multiple_files=True, key="context_uploader")
        with col_adv_2:
            st.markdown("**雙向細目表單元設定**")
            if context_files:
                unit_count = st.number_input("單元數量", min_value=1, max_value=10, value=3)
                cols = st.columns(unit_count)
                for i in range(unit_count):
                    with cols[i]:
                        u_name = st.text_input(f"單元 {i+1}", key=f"unit_{i}")
                        if u_name: unit_list.append(u_name)
            else:
                st.info("💡 請先上傳教材以啟用單元編輯")

    with st.expander("🧠 AI 模型核心設定", expanded=False):
        model_mode = st.radio("模式", ["智慧分流 (建議)", "手動指定"], label_visibility="collapsed")
        if model_mode == "手動指定":
            manual_model = st.selectbox("核心", ["Gemini 3.0 Pro", "Gemini 3.0 Flash"], label_visibility="collapsed")

# --- 審查標準 (隱藏) ---
LITERACY_STANDARDS = """
檢核標準：試著將題目中的「情境敘述」（故事、圖片、前言）移除，並依下方4個項度及各科真假素養審查標準檢核。
1.情境真實性：判斷情境是否真實、是否貼近學生生活經驗。
2.推理需求：分析學生是否需要進行推論、判斷或解釋，而非直接記憶。
3.跨概念整合：檢查是否整合多個概念或學習單元。
4.情境包裝程度：判斷情境是否只是包裝知識，而非解題必要條件。

【各科真假素養審查標準】：
1. 國語科：(真)閱讀依存、高階思維、多元表徵；(假)情境脫節、低階提問。
2. 數學科：(真)功能性情境、真實解題(含雜訊)、數學建模；(假)文字堆砌、數據完美、套路解題。
3. 英語科：(真)真實語料、語用溝通、資訊素養；(假)去脈絡化、死記硬背、文化真空。
4. 社會科：(真)史料判讀、多重觀點、因果探究；(假)瑣碎記憶、單一觀點、結論背誦。
5. 自然科：(真)探究歷程、解釋現象、證據論述；(假)名詞解釋、結果背誦、違背常理。
6. 生活課程：(真)感官體驗、情境應變、實作導向；(假)規訓教條、知識超載、文字負擔。
"""

st.markdown("---")

if start_btn:
    if not exam_file:
        st.warning("❌ 請務必上傳一份「試卷」！")
    else:
        status_box = st.empty()
        progress_bar = st.progress(0)
        
        try:
            # ----------------------------------------------------
            # Phase 1: 智慧分流路由 (Smart Routing)
            # ----------------------------------------------------
            filename = exam_file.name
            status_box.info(f"🔍 正在解析試卷資訊... (檔名：{filename})")
            progress_bar.progress(10)

            # 1. 判斷科目屬性
            is_science = False
            if any(k in filename for k in ["數學", "自然", "理化", "物理", "化學", "生物"]):
                is_science = True
            
            # 2. 決定模型 (核心路由邏輯)
            target_model_name = ""
            routing_msg = ""

            if model_mode == "手動指定" and ENABLE_ADVANCED_FEATURES:
                if "Pro" in manual_model:
                    target_model_name = get_best_pro_model(api_key)
                    routing_msg = "🧠 手動指定 Pro 模型"
                else:
                    target_model_name = get_best_flash_model(api_key)
                    routing_msg = "⚡ 手動指定 Flash 模型"
            else:
                # 恢復品質優先的智慧分流
                if context_files:
                    target_model_name = get_best_pro_model(api_key)
                    routing_msg = "📚 偵測到參考教材，啟用品質優先分析模式"
                elif is_science:
                    target_model_name = get_best_pro_model(api_key)
                    routing_msg = "📐 理科試卷分析"
                else:
                    target_model_name = get_best_flash_model(api_key)
                    routing_msg = "📝 文科試卷分析"

            status_box.info(f"🔄 AI 審查中 ... ")

            # 3. 資訊提取 (Metadata)
            flash_model = genai.GenerativeModel(get_best_flash_model(api_key))
            exam_ref = upload_to_gemini(exam_file)
            
            meta_prompt = """
            請閱讀這份試卷，並擷取以下資訊，輸出為純 JSON 格式：
            {
                "year": "學年度 (例如 113)",
                "semester": "學期 (例如 第一學期)",
                "grade": "年級 (例如 六年級)",
                "subject": "科目 (例如 數學)",
                "exam_type": "評量類別 (例如 期末評量)"
            }
            如果找不到某些資訊，請填入空白。
            """
            meta_response = flash_model.generate_content([meta_prompt, exam_ref])
            try:
                json_str = meta_response.text.strip()
                if "```json" in json_str:
                    json_str = json_str.split("```json")[1].split("```")[0]
                metadata = json.loads(json_str)
            except:
                metadata = {"year":"", "semester":"", "grade":"", "subject":"", "exam_type":""}
            st.session_state.metadata = metadata

            # ----------------------------------------------------
            # Phase 2: 深度審查 (Updated Prompt v5.6 - Final Refined)
            # ----------------------------------------------------
            
            # 設定生成參數：Temperature = 0 (強制理性)
            generation_config = {
                "temperature": 0.0,
                "top_p": 1.0,
                "top_k": 32,
            }
            
            main_model = genai.GenerativeModel(target_model_name)
            
            prompt_parts = []
            if context_files:
                for cf in context_files: prompt_parts.append(upload_to_gemini(cf))
            
            units_str = ", ".join(unit_list) if unit_list else "未提供"

            base_prompt = f"""
你是一位精通「台灣 108 課綱素養導向評量」的試題審查專家。
目前正在審查：{metadata.get('year')}學年度 {metadata.get('subject')} 試卷。

**【台灣試卷三大排版閱讀協定 (Taiwan Exam Layout Protocol)】：**
請先掃描整份試卷的幾何結構，並嚴格依照下列 **3 種模式** 擇一執行閱讀：

**Mode 1: 不分欄 (Single Column)**
* **特徵**：A4或A3版面，整頁無明顯分隔線。
* **閱讀順序**：標準 Z 字型（由左至右 ➡，由上而下 ⬇）。

**Mode 2: 左右雙欄 (Left-Right Split)**
* **特徵**：B4版面，中間有一條「垂直分隔線」，文字橫書（常見於數學、自然、社會、英文）。
* **閱讀順序**：
    1.  **先讀左欄**：在左半頁範圍內，由左至右、由上而下讀取。
    2.  **再讀右欄**：移至右半頁，由左至右、由上而下讀取。
    * *注意：嚴禁跨欄閱讀（即不可直接從左欄第一行讀到右欄第一行）。*

**Mode 3: 上下分欄 (Top-Bottom Split)**
* **特徵**：B4版面，中間有一條「水平分隔線」，文字為**直書**（由上而下排列，常見於國語文）。
* **閱讀順序**：
    1.  **先讀上欄**：在上半頁範圍內，**由右至左 ⬅** 掃描直行文字。
    2.  **再讀下欄**：移至下半頁，**由右至左 ⬅** 掃描直行文字。

**2. 跨欄/跨頁拼接技術 (Cross-Page Stitching)**：
* **「題號」是你的唯一導航**：在閱讀時，請隨時追蹤題號（1, 2, 3...）。
* **斷點偵測**：當一個欄位或頁面結束時，若最後一題（例如第 5 題）只有題幹沒有選項，**請立即去「下一欄頂端」或「下一頁開頭」尋找該題的剩餘部分**。
* **嚴禁幻覺**：如果找不到接續的內容，請標記該題為「題目內容不完整」，絕對不可自行編造選項。

**【台灣國小教學情境守則 (Contextual Rules)】：**
1. **是非題/改錯題 絕對豁免**：在審查「是非題」或「改錯題」時，若題目敘述的錯誤是為了測驗學生觀念（且標準答案為「X」或「不正確」），這是**正確的命題設計**。請將其歸類為 **優良試題**，**嚴禁**列為待改善或假素養。只有當「錯誤觀念」被標示為「正確答案 (O)」時，才視為命題錯誤。
2. **在地化教學標準**：針對讀音、定義或格式，以台灣國小教學現場慣例為準。

**【排版與精準度憲法】：**
1. 嚴禁開場白。
2. 禁止頁碼，精準對應題號。
3. 強制換行。
4. **題號格式統一**：請務必使用「**大題-小題**」格式 (例如：**二-7**、**三-1**)，嚴禁使用「題二-第7題」這種冗贅寫法。
5. **題目錨點**：務必於題號後方，括號摘錄該題題目開頭約 5~7 個字。
6. **拒絕模糊論述**：每一項分析都必須具體列出是哪幾題。
7. **先在內部完成檢查，再輸出最終答案**：請自行檢查題號是否前後一致、表格是否完整、百分比是否合理後，再輸出。
8. **除指定表格外，禁止自行發明其他表格或格式。**

請嚴格依照以下順序輸出 Markdown 報告：

## 總結與建議
### 🔴 最優先修正 (Critical)
若存在重大錯誤（如答案錯誤、題意錯誤、無法作答題目），請列出並說明；若無，請寫「無重大錯誤」。

### ⚖️ 難度與鑑別度點評
整體評估試卷難度分布與鑑別度是否合理。

### 👍 值得讚許之處
列出試卷優點（最多 5 點即可）。

### 💡 後續優化建議
提出可提升試卷品質的建議。

## 題幹與邏輯品質
評估每一題時，請從以下三個面向進行內部檢查：
1. 題幹完整性(檢查題幹資訊是否完整、是否缺乏必要條件、是否存在無法作答情形。)
2. 選項與干擾品質(檢查選項是否具有合理干擾性，是否存在明顯錯誤選項或選項長度差異過大。)
3. 是否存在提示線索(檢查題幹或選項是否提供過度提示，使學生可不經思考直接猜出答案。)

輸出時不必逐項列出檢查結果，而是在每一題的說明中自然說明其優點或問題。

### 🟢 優良試題
列出具有代表性的優良題目即可（3～5題）。  
格式：
* 一-2（題目開頭）— 題幹清楚、選項干擾合理，能有效檢核學生概念理解。
* 三-1（題目開頭）— 題意明確，作答需理解情境而非直接記憶。
* (其餘優良試題略)...
(**強制清單**：每一題務必換行，使用列點符號 (*) 開頭，每一點都必須包含「題號 + 題目錨點 + 具體說明原因。)

### 🟡 待確認試題
必須**完整列出所有有問題的題目**，不可省略。  
每題需具體說明問題，例如題幹不完整、干擾選項不足、存在提示線索等。

* 四-2（題目開頭）— 題幹資訊不足，未說明條件，可能造成多重解讀。
* 五-1（題目開頭）— 選項 C 與 D 明顯錯誤，干擾效果不足。
(**強制清單**：每一題務必換行，使用列點符號 (*) 開頭，每一點都必須包含「題號 + 題目錨點 + 具體說明原因。)

* **⚠️ 強制豁免守則**：
    * 是非題/改錯題/選錯題：若錯誤敘述對應標準答案為「X」或「選出錯誤選項」，視為 **🟢 優良試題**。

## 素養導向深度審查
{LITERACY_STANDARDS}
### 🟢 真素養
列出代表性題目（3～5題），並簡述其情境與能力要求。
(**強制清單**：每一題務必換行，使用列點符號 (*) 開頭，每一點都必須包含「題號 + 題目錨點 + 具體說明原因。)

### 🟡 假素養/待確認
**完整列出所有問題題目**，說明原因，例如：
- 情境只是裝飾
- 不需要推理即可作答
- 未整合概念
(**強制清單**：每一題務必換行，使用列點符號 (*) 開頭，每一點都必須包含「題號 + 題目錨點 + 具體說明原因。)

## 公平性與敏感度審查

評估每一題時，請從以下兩個面向進行內部檢查：
1.文化公平：檢查是否涉及文化偏見或刻板印象（如性別、種族、語文、文化、職業等）。
2.情境熟悉度：檢查題目情境是否可能造成背景知識差異。

### 🟢 通過 (無敏感議題)
列出代表性題目即可（1～3題），最後可加：

* (其餘通過試題略)...
(**強制清單**：每一題務必換行，使用列點符號 (*) 開頭，每一點都必須包含「題號 + 題目錨點 + 具體說明原因。)

### 🟡 潛在爭議 (具體指出問題)
若存在文化偏見、情境背景差異等問題，需**完整列出所有相關題目**。
(**強制清單**：每一題務必換行，使用列點符號 (*) 開頭，每一點都必須包含「題號 + 題目錨點 + 具體說明原因。)

## 雙向細目表核算
請**只能**輸出以下 Markdown 表格，不可改標題、不可加前言、不可加註解、不可少列：

| 認知向度 | 對應題號 | 比重 |
| --- | --- | --- |
| 記憶 |  |  |
| 理解 |  |  |
| 應用 |  |  |
| 分析 |  |  |
| 評鑑 |  |  |
| 創造 |  |  |

硬性規定：
1. 必須剛好六列，不可增減。
2. 「比重」一律用百分比表示，例如 15%。
3. 六列百分比總和必須為 100%。
4. 若無法完全判定，也要依最合理方式分配，不可留空整列。
5. 本段除了上表，不可輸出任何額外句子。

## 難易度與負擔分析
請**只能**輸出以下 Markdown 表格，不可改標題、不可加前言、不可加註解：

| 難易度 | 對應題號 | 比重 |
| --- | --- | --- |
| 易 |  |  |
| 中 |  |  |
| 難 |  |  |

硬性規定：
1. 必須剛好三列，不可增減。
2. 「比重」一律用百分比表示，例如 40%。
3. 三列百分比總和必須為 100%。
4. 「對應題號」請寫具體題號，不可只寫「略」或「多題」。

接著輸出整體作答負擔觀察：

### 📖 閱讀負擔  
簡述試卷整體閱讀量是否適中，是否有長題幹或大量閱讀題。

### 📊 圖表判讀負擔  
分析是否需要解讀圖表、數據或圖像資訊。

### 🧮 運算負擔  
分析是否需要多步驟計算或複雜推理。

若整體負擔適中，可簡要說明即可，不需過度分析。

"""
            if context_files: prompt_parts.append("【參考教材】：")
            prompt_parts.append(base_prompt)
            prompt_parts.append("【待審查試卷】：")
            prompt_parts.append(exam_ref)
            
            progress_bar.progress(60)
            
            # 帶入 generation_config
            response = main_model.generate_content(
                prompt_parts,
                generation_config=generation_config
            )
            
            progress_bar.progress(100)
            
            # 1. 先處理換行
            raw_text = response.text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
            
            # 2. 呼叫清洗函式
            final_cleaned_text = clean_ai_hallucinations(raw_text)
            
            status_box.success(f"✅ 分析完成！")
            
            st.session_state.analysis_result = final_cleaned_text
            st.session_state.used_model_name = target_model_name
            
        except Exception as e:
            st.error(f"發生錯誤: {e}")
            if "429" in str(e): st.warning("💡 提示：目前 AI 忙線中，請稍後再試。")

# --- 結果顯示與 Word 生成 ---
if st.session_state.analysis_result:
    normalized_result = normalize_analysis_tables(st.session_state.analysis_result)

    st.warning("⚠️ **系統限制與聲明**：本系統僅針對試題內容進行深度分析，**未檢核**「命題範圍」與「試卷形式」（如題號連貫性、配分加總正確性），請老師務必自行審閱。")

    if "## 題幹與邏輯品質" in normalized_result:
        summary_part, body_part = normalized_result.split("## 題幹與邏輯品質", 1)
        body_part = "## 題幹與邏輯品質" + body_part
        
        summary_part = re.sub(r'^\s*[\*\-]\s+(###)', r'\1', summary_part, flags=re.MULTILINE)
        summary_part = re.sub(r'\n\s*---\s*$', '', summary_part).strip()

        st.info(summary_part)
        st.markdown("---")
        st.markdown(body_part)
    else:
        st.markdown("## 📊 審查報告")
        st.markdown(normalized_result)
    
    word_binary = create_word_report(normalized_result, st.session_state.metadata)
    
    st.download_button(
        label="📥 下載 Word 報告 (.docx)",
        data=word_binary,
        file_name=f"建功國小_{st.session_state.metadata.get('subject','科目')}_審題報告.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_btn"
    )

render_footer()
